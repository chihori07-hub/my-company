from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# フォント設定
style = doc.styles['Normal']
style.font.name = 'Hiragino Kaku Gothic ProN'
style.font.size = Pt(11)

# タイトル
title = doc.add_heading('ローンチセミナー 内容設計案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('作成日：2026年5月7日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# 基本情報
doc.add_heading('基本情報', 1)
info = doc.add_paragraph()
info.add_run('対象サービス：').bold = True
info.add_run('夫婦関係を再構築する7つの実践プログラム')
doc.add_paragraph()
info2 = doc.add_paragraph()
info2.add_run('形式：').bold = True
info2.add_run('60分・無料・オンライン')
doc.add_paragraph()
info3 = doc.add_paragraph()
info3.add_run('導線：').bold = True
info3.add_run('セミナー → 個別相談 → プログラム申し込み')

doc.add_paragraph('')

# ターゲット
doc.add_heading('ターゲット', 1)
doc.add_paragraph('「夫を変えようとしてきたけど、それじゃないと薄々感じている」30代後半〜50代前半の女性。\nすでに本やセミナーで「自分を見直そう」という情報には触れてきたが、現実が変わっていない人。')

doc.add_paragraph('')

# 差別化ポイント
doc.add_heading('差別化ポイント', 1)
p1 = doc.add_paragraph()
p1.add_run('SNSに多い発信：').bold = True
p1.add_run('「夫のせいじゃない、自分の本音に気づくことが大事」')

doc.add_paragraph('')

p2 = doc.add_paragraph()
p2.add_run('このセミナーが伝えること：').bold = True

doc.add_paragraph('本音に気づこうとしても気づけないのは、思い込み・セルフイメージ（＝前提）が本音を覆っているから。\n前提が変わらないまま本音を探しても、出てくるのは「前提にくっついた気持ち」であって、本当の本音ではない。')

p3 = doc.add_paragraph()
p3.add_run('だから、先に前提を見つめ直す。').bold = True

doc.add_paragraph('')

# テーマ案
doc.add_heading('テーマ案（要ブラッシュアップ）', 1)
theme = doc.add_paragraph()
run = theme.add_run('「本音に気づこうとしても、できなかった理由」\n——思い込みが変わると、本当の気持ちが動き出す')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph('')

# 60分構成
doc.add_heading('60分 構成案', 1)

# ①
doc.add_heading('① つかみ・共感　10分', 2)
doc.add_paragraph('「私ばかり頑張っているのに、夫が動いてくれない」という感覚を言語化。\n参加者が「私のことだ」と感じる入口を作る。')

doc.add_paragraph('')

# ②
doc.add_heading('② 視点の転換体験　20分　※セミナーの核心', 2)
doc.add_paragraph('セッションの層の構造を図で見せながら、ゆみさんのストーリーで具体的に説明する。')

doc.add_paragraph('')

# 層の構造
p_layer = doc.add_paragraph()
p_layer.add_run('【セッションで掘り下げる層の構造】').bold = True

layers = [
    '出来事（何が起きたか）',
    '思考の言語化（その時、頭の中にどんな言葉が流れているか）',
    '感情の表層（その時、どんな感情を感じているか）',
    '感情の深層（さらに奥にある感情は何か）',
    '心の欲求（本当はどうなりたいか・どんな自分でありたいか）',
]
for i, layer in enumerate(layers):
    p = doc.add_paragraph(layer, style='List Number')

doc.add_paragraph('')

p_diff = doc.add_paragraph()
p_diff.add_run('SNSは感情の表層で止まっている。このアプローチは深層まで降りていく。').bold = True
doc.add_paragraph('前提（思い込み）は深層と心の欲求の間に埋まっているから、表層だけ見ても変わらない。')

doc.add_paragraph('')

# ゆみさんのストーリー
p_story = doc.add_paragraph()
p_story.add_run('【ゆみさんのストーリー（語り方の流れ）】').bold = True

story_steps = [
    '夫への不満を語るゆみさん（「私ばかり頑張っている、夫が動いてくれない」）',
    'コーチがゆみさんの言葉の中から特定の場面を見つけ、その文脈の中で気遣いを肯定する',
    '肯定した上で視点を反転させる問いを投げる\n　→「ゆみさんが察して先回りしてやってしまうことで、旦那様がそれをするタイミングを逃しているということはありませんか？」',
    'ゆみさんの反応「あぁ……それはあるかもしれません」',
    '「これは、あなたが悪いということではない。『私がやらなければ』という前提が、無意識にそうさせていた」',
]
for step in story_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph('')

# ③
doc.add_heading('③ 前提とは何か　10分', 2)
items = [
    '前提＝思い込み・セルフイメージ',
    '前提が行動・感情・関係性を作っている',
    '前提が変わらないと、行動を変えようとしても元に戻る',
    '本音に気づけないのも、前提が本音を覆っているから',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ④
doc.add_heading('④ プログラム紹介　10分', 2)
items4 = [
    '7つの実践プログラムで何をするか',
    'どんな変化が起きるか',
    'どんな方に向いているか',
]
for item in items4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ⑤
doc.add_heading('⑤ 個別相談へ誘導　10分', 2)
doc.add_paragraph('「今日お話しした層の掘り下げは、自分一人ではなかなかできません。コーチの問いかけがあって初めて、自分では気づけなかった前提に気づける。個別相談では、あなた自身の前提を一緒に見ていきます。」')

doc.add_paragraph('')

# 検討事項
doc.add_heading('検討事項（コンサルに相談したいこと）', 1)
items_check = [
    'テーマ・タイトルの言葉のブラッシュアップ',
    '②パートの「層の構造図」をどう見せるか（スライド設計）',
    'ゆみさんのストーリーの語り方・長さのバランス',
    'プログラム紹介の内容と順番',
    '個別相談への誘導の言葉',
]
for item in items_check:
    doc.add_paragraph('□ ' + item)

doc.add_paragraph('')
doc.add_paragraph('Coa-Terasu　東山知帆里').alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save('/Users/higashiyamachihori/my-company/ローンチセミナー内容設計案.docx')
print('完了')
